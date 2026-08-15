import os
import sqlite3
import re
import io
import html as html_lib
import secrets

from flask import (
    Flask,
    render_template,
    request,
    redirect,
    session,
    flash,
    send_file,
    send_from_directory
)

from werkzeug.security import (
    generate_password_hash,
    check_password_hash
)

from werkzeug.utils import secure_filename

import markdown
import base64

from groq import Groq

from PyPDF2 import PdfReader

import pytesseract
from pdf2image import convert_from_path

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from docx import Document
from docx.shared import Pt


# =========================================================
# FLASK CONFIGURATION
# =========================================================

app = Flask(__name__)

app.secret_key = "askai_super_secret_key_change_this"


# =========================================================
# GROQ AI SETUP (fast cloud AI, replaces local Ollama)
# =========================================================

def load_groq_api_key():

    env_key = os.environ.get("GROQ_API_KEY")

    if env_key:

        return env_key.strip()


    key_file_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "groq_key.txt"
    )

    if os.path.exists(key_file_path):

        with open(key_file_path, "r", encoding="utf-8") as f:

            return f.read().strip()


    return None


GROQ_API_KEY = load_groq_api_key()

groq_client = (
    Groq(api_key=GROQ_API_KEY)
    if GROQ_API_KEY
    else None
)

GROQ_TEXT_MODEL = "openai/gpt-oss-20b"

GROQ_VISION_MODEL = "qwen/qwen3.6-27b"


def groq_not_configured_message():

    return (
        "⚠️ Groq API key not set up. Create a file named "
        "'groq_key.txt' in your project folder and paste your "
        "Groq API key inside it (get a free key at "
        "console.groq.com)."
    )


UPLOAD_FOLDER = "uploads"

ALLOWED_EXTENSIONS = {"pdf"}

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# =========================================================
# DATABASE
# =========================================================

DATABASE = "database.db"


def get_db():

    conn = sqlite3.connect(DATABASE)

    conn.row_factory = sqlite3.Row

    return conn


def create_database():

    conn = get_db()

    cursor = conn.cursor()


    # USERS

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL
        )
    """)


    # CHAT HISTORY

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS chat_history(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            question TEXT NOT NULL,
            answer TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)


    # PDF TABLE

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS pdf_documents(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            pdf_name TEXT NOT NULL,
            text_file TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)


    conn.commit()


    # MIGRATION: add pdf_file column if it does not already exist
    # (needed for multiple-PDF support)

    try:

        cursor.execute(
            "ALTER TABLE pdf_documents ADD COLUMN pdf_file TEXT"
        )

        conn.commit()

    except sqlite3.OperationalError:

        pass


    # MIGRATION: add image_file column if it does not already exist
    # (needed for image-attach chat support)

    try:

        cursor.execute(
            "ALTER TABLE chat_history ADD COLUMN image_file TEXT"
        )

        conn.commit()

    except sqlite3.OperationalError:

        pass


    # MIGRATION: add rating column if it does not already exist
    # (needed for 👍/👎 answer feedback)

    try:

        cursor.execute(
            "ALTER TABLE chat_history ADD COLUMN rating INTEGER"
        )

        conn.commit()

    except sqlite3.OperationalError:

        pass


    # MIGRATION: add share_id column if it does not already exist
    # (needed for shareable chat links)

    try:

        cursor.execute(
            "ALTER TABLE chat_history ADD COLUMN share_id TEXT"
        )

        conn.commit()

    except sqlite3.OperationalError:

        pass


    conn.close()


create_database()


# =========================================================
# HELPER FUNCTIONS
# =========================================================

def allowed_file(filename):

    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower()
        in ALLOWED_EXTENSIONS
    )


ALLOWED_IMAGE_EXTENSIONS = {
    "png",
    "jpg",
    "jpeg",
    "gif",
    "webp"
}


def allowed_image_file(filename):

    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower()
        in ALLOWED_IMAGE_EXTENSIONS
    )


def clean_text(text):

    if not text:
        return ""

    text = re.sub(
        r"\s+",
        " ",
        text
    )

    return text.strip()


def get_user_pdfs():

    if "user" not in session:
        return []

    conn = get_db()

    cursor = conn.cursor()

    cursor.execute("""
        SELECT *
        FROM pdf_documents
        WHERE username=?
        ORDER BY id DESC
    """, (session["user"],))

    pdfs = cursor.fetchall()

    conn.close()

    return pdfs


def get_active_pdf():

    if "user" not in session:
        return None

    pdfs = get_user_pdfs()

    if not pdfs:

        session.pop("active_pdf_id", None)

        return None


    active_id = session.get("active_pdf_id")

    if active_id:

        for pdf in pdfs:

            if pdf["id"] == active_id:

                return pdf


    # No active PDF set (or it was removed) -> fall back to most recent

    session["active_pdf_id"] = pdfs[0]["id"]

    return pdfs[0]


# =========================================================
# EXPORT HELPERS (Answer -> PDF / Word)
# =========================================================

def html_answer_to_text(answer_html):

    if not answer_html:
        return ""


    # Turn common block tags into line breaks before stripping

    text = re.sub(
        r"</(p|div|li|h[1-6])>",
        "\n",
        answer_html,
        flags=re.IGNORECASE
    )

    text = re.sub(
        r"<br\s*/?>",
        "\n",
        text,
        flags=re.IGNORECASE
    )


    # Strip all remaining HTML tags

    text = re.sub(
        r"<[^>]+>",
        "",
        text
    )


    text = html_lib.unescape(text)


    # Collapse extra blank lines

    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text
    )

    return text.strip()


def generate_pdf_file(question, answer_html):

    from xml.sax.saxutils import escape as xml_escape

    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch
    )


    styles = getSampleStyleSheet()

    story = []


    story.append(
        Paragraph(
            "AskAI - Question & Answer",
            styles["Title"]
        )
    )

    story.append(Spacer(1, 14))


    story.append(
        Paragraph(
            "<b>Question:</b> " + xml_escape(question),
            styles["Heading3"]
        )
    )

    story.append(Spacer(1, 10))


    answer_text = html_answer_to_text(answer_html)

    for paragraph in answer_text.split("\n\n"):

        if not paragraph.strip():
            continue

        safe_paragraph = xml_escape(
            paragraph
        ).replace("\n", "<br/>")

        story.append(
            Paragraph(
                safe_paragraph,
                styles["BodyText"]
            )
        )

        story.append(Spacer(1, 8))


    doc.build(story)

    buffer.seek(0)

    return buffer


def generate_docx_file(question, answer_html):

    buffer = io.BytesIO()

    document = Document()


    document.add_heading(
        "AskAI - Question & Answer",
        level=1
    )


    document.add_heading(
        "Question",
        level=2
    )

    q_paragraph = document.add_paragraph(question)

    q_paragraph.runs[0].font.size = Pt(12)


    document.add_heading(
        "Answer",
        level=2
    )


    answer_text = html_answer_to_text(answer_html)

    for paragraph in answer_text.split("\n\n"):

        if not paragraph.strip():
            continue

        document.add_paragraph(paragraph)


    document.save(buffer)

    buffer.seek(0)

    return buffer


# =========================================================
# PDF TEXT EXTRACTION
# =========================================================

def extract_pdf_text(pdf_path):

    extracted_text = ""


    # -----------------------------------------------------
    # NORMAL PDF TEXT EXTRACTION
    # -----------------------------------------------------

    try:

        reader = PdfReader(pdf_path)

        for page in reader.pages:

            try:

                page_text = page.extract_text()

                if page_text:

                    extracted_text += "\n" + page_text

            except Exception as e:

                print(
                    "PAGE TEXT ERROR:",
                    e
                )

    except Exception as e:

        print(
            "PDF READER ERROR:",
            e
        )


    extracted_text = clean_text(
        extracted_text
    )


    # -----------------------------------------------------
    # OCR FALLBACK
    # -----------------------------------------------------

    if len(extracted_text) < 50:

        print(
            "No readable text found."
        )

        print(
            "Starting OCR..."
        )

        try:

            images = convert_from_path(
                pdf_path,
                dpi=200
            )


            ocr_text = ""


            for index, image in enumerate(images):

                print(
                    f"OCR page {index + 1}"
                )

                text = pytesseract.image_to_string(
                    image
                )

                ocr_text += "\n" + text


            extracted_text = clean_text(
                ocr_text
            )


        except Exception as e:

            print(
                "OCR ERROR:",
                e
            )


    return extracted_text


# =========================================================
# FIND RELEVANT PDF TEXT
# =========================================================

def find_relevant_text(
    full_text,
    question,
    max_chars=12000
):

    if not full_text:

        return ""


    # Split into chunks

    chunks = re.split(
        r"(?<=[.!?])\s+",
        full_text
    )


    words = re.findall(
        r"\b[a-zA-Z]{3,}\b",
        question.lower()
    )


    stop_words = {
        "what",
        "when",
        "where",
        "which",
        "who",
        "whom",
        "this",
        "that",
        "from",
        "with",
        "about",
        "explain",
        "tell",
        "give",
        "does",
        "the",
        "and",
        "are",
        "was",
        "were",
        "how",
        "why"
    }


    keywords = [
        word
        for word in words
        if word not in stop_words
    ]


    scored_chunks = []


    for chunk in chunks:

        lower_chunk = chunk.lower()

        score = 0


        for keyword in keywords:

            if keyword in lower_chunk:

                score += 1


        if score > 0:

            scored_chunks.append(
                (
                    score,
                    chunk
                )
            )


    scored_chunks.sort(
        key=lambda x: x[0],
        reverse=True
    )


    selected = []


    current_length = 0


    for score, chunk in scored_chunks:

        if current_length + len(chunk) > max_chars:

            break


        selected.append(chunk)

        current_length += len(chunk)


    # If no keyword match

    if not selected:

        return full_text[:max_chars]


    return " ".join(selected)


# =========================================================
# AI NORMAL CHAT
# =========================================================

def ask_ai(question):

    if not groq_client:

        return groq_not_configured_message()


    try:

        response = groq_client.chat.completions.create(

            model=GROQ_TEXT_MODEL,

            messages=[
                {
                    "role": "user",
                    "content": question
                }
            ]

        )


        answer = response.choices[0].message.content


        return answer


    except Exception as e:

        print(
            "GROQ ERROR:",
            e
        )

        return (
            "AI Error: "
            + str(e)
        )


# =========================================================
# AI IMAGE QUESTION ANSWERING
# =========================================================

def ask_image_ai(question, image_path):

    if not groq_client:

        return groq_not_configured_message()


    try:

        ext = image_path.rsplit(".", 1)[-1].lower()

        mime_type = {
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "gif": "image/gif",
            "webp": "image/webp"
        }.get(ext, "image/jpeg")


        with open(image_path, "rb") as image_file:

            encoded_image = base64.b64encode(
                image_file.read()
            ).decode("utf-8")


        data_url = (
            "data:" + mime_type + ";base64," + encoded_image
        )


        response = groq_client.chat.completions.create(

            model=GROQ_VISION_MODEL,

            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                question
                                if question
                                else "Describe this image in detail."
                            )
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": data_url
                            }
                        }
                    ]
                }
            ]

        )


        answer = response.choices[0].message.content


        return answer


    except Exception as e:

        print(
            "IMAGE AI ERROR:",
            e
        )

        return (
            "AI Error while analyzing the image: "
            + str(e)
        )


# =========================================================
# SUGGESTED FOLLOW-UP QUESTIONS
# =========================================================

# =========================================================
# AI PDF QUESTION ANSWERING
# =========================================================

def ask_pdf_ai(
    question,
    pdf_context
):

    prompt = f"""
You are a professional PDF Question Answering Assistant.

Answer the user's question ONLY using the PDF context provided below.

If the answer is not available in the PDF context, clearly say:

"The answer was not found in the uploaded PDF."

Do not invent facts.

Give a clear and useful answer.

PDF CONTEXT:
-------------------------
{pdf_context}
-------------------------

USER QUESTION:
{question}

ANSWER:
"""


    if not groq_client:

        return groq_not_configured_message()


    try:

        response = groq_client.chat.completions.create(

            model=GROQ_TEXT_MODEL,

            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]

        )


        answer = response.choices[0].message.content


        return answer


    except Exception as e:

        print(
            "PDF AI ERROR:",
            e
        )

        return (
            "AI Error: "
            + str(e)
        )


# =========================================================
# PWA SERVICE WORKER (served at root so it can control the
# whole site, not just /static/)
# =========================================================

@app.route("/sw.js")
def service_worker():

    return send_from_directory(
        "static",
        "sw.js",
        mimetype="application/javascript"
    )


# =========================================================
# HOME
# =========================================================

@app.route("/")
def home():

    if "user" not in session:

        return redirect("/login")

    return redirect("/chat")


# =========================================================
# GUEST CHATBOT (no login required)
# =========================================================

@app.route(
    "/bot",
    methods=["GET", "POST"]
)
def guest_bot():

    answer = ""

    question = ""


    if request.method == "POST":

        question = request.form.get(
            "question",
            ""
        ).strip()


        if question:

            raw_answer = ask_ai(question)

            answer = markdown.markdown(raw_answer)

        else:

            flash(
                "Please enter a question.",
                "error"
            )


    return render_template(
        "guest_chat.html",
        answer=answer,
        question=question
    )


# =========================================================
# LOGIN
# =========================================================

@app.route(
    "/login",
    methods=["GET", "POST"]
)
def login():

    if request.method == "POST":

        email = request.form.get(
            "email",
            ""
        ).strip()


        password = request.form.get(
            "password",
            ""
        )


        conn = get_db()

        cursor = conn.cursor()


        cursor.execute("""
            SELECT *
            FROM users
            WHERE email=?
        """, (email,))


        user = cursor.fetchone()

        conn.close()


        if user and check_password_hash(
            user["password"],
            password
        ):

            session["user"] = user["username"]

            flash(
                "Login successful!",
                "success"
            )

            return redirect("/dashboard")


        flash(
            "Invalid email or password.",
            "error"
        )


    return render_template(
        "login.html"
    )


# =========================================================
# SIGNUP
# =========================================================

@app.route(
    "/signup",
    methods=["GET", "POST"]
)
def signup():

    if request.method == "POST":

        username = request.form.get(
            "username",
            ""
        ).strip()


        email = request.form.get(
            "email",
            ""
        ).strip()


        password = request.form.get(
            "password",
            ""
        )


        if not username or not email or not password:

            flash(
                "Please fill all fields.",
                "error"
            )

            return redirect("/signup")


        hashed_password = generate_password_hash(
            password
        )


        conn = get_db()

        cursor = conn.cursor()


        try:

            cursor.execute(
                """
                INSERT INTO users(
                    username,
                    email,
                    password
                )
                VALUES(?,?,?)
                """,
                (
                    username,
                    email,
                    hashed_password
                )
            )


            conn.commit()

            flash(
                "Account created successfully!",
                "success"
            )


            return redirect("/login")


        except sqlite3.IntegrityError:

            flash(
                "Email already exists.",
                "error"
            )


        finally:

            conn.close()


    return render_template(
        "signup.html"
    )


# =========================================================
# LOGOUT
# =========================================================

@app.route("/logout")
def logout():

    session.clear()

    flash(
        "You have been logged out.",
        "success"
    )

    return redirect("/login")


# =========================================================
# DASHBOARD
# =========================================================

@app.route("/dashboard")
def dashboard():

    if "user" not in session:

        return redirect("/login")


    username = session["user"]


    conn = get_db()

    cursor = conn.cursor()


    # Questions

    cursor.execute("""
        SELECT COUNT(*)
        FROM chat_history
        WHERE username=?
    """, (username,))


    total_questions = cursor.fetchone()[0]


    # PDFs

    cursor.execute("""
        SELECT COUNT(*)
        FROM pdf_documents
        WHERE username=?
    """, (username,))


    total_pdfs = cursor.fetchone()[0]


    # Recent chats

    cursor.execute("""
        SELECT *
        FROM chat_history
        WHERE username=?
        ORDER BY id DESC
        LIMIT 5
    """, (username,))


    recent_chats = cursor.fetchall()


    conn.close()


    return render_template(
        "dashboard.html",
        username=username,
        total_questions=total_questions,
        total_pdfs=total_pdfs,
        recent_chats=recent_chats
    )


# =========================================================
# SERVE UPLOADED IMAGES
# =========================================================

@app.route("/uploads/<path:filename>")
def uploaded_file(filename):

    if "user" not in session:

        return redirect("/login")

    return send_from_directory(
        app.config["UPLOAD_FOLDER"],
        filename
    )


# =========================================================
# CHAT
# =========================================================

@app.route(
    "/chat",
    methods=["GET", "POST"]
)
def chat():

    if "user" not in session:

        return redirect("/login")


    answer = ""

    question = ""


    if request.method == "POST":

        question = request.form.get(
            "question",
            ""
        ).strip()


        image_file_obj = request.files.get("image")

        has_image = (
            image_file_obj
            and image_file_obj.filename != ""
        )


        if not question and not has_image:

            flash(
                "Please enter a question or attach an image.",
                "error"
            )

            return redirect("/chat")


        image_saved_name = None


        try:

            # ------------------------------------------------
            # IF IMAGE ATTACHED -> vision AI answers
            # ------------------------------------------------

            if has_image:

                if not allowed_image_file(
                    image_file_obj.filename
                ):

                    flash(
                        "Only image files (png, jpg, jpeg, gif, webp) "
                        "are allowed.",
                        "error"
                    )

                    return redirect("/chat")


                username = secure_filename(
                    session["user"]
                )

                original_image_name = secure_filename(
                    image_file_obj.filename
                )

                image_saved_name = (
                    username
                    + "_img_"
                    + str(int(__import__("time").time()))
                    + "_"
                    + original_image_name
                )

                image_path = os.path.join(
                    app.config["UPLOAD_FOLDER"],
                    image_saved_name
                )

                image_file_obj.save(image_path)


                answer = ask_image_ai(
                    question,
                    image_path
                )


                if not question:

                    question = "🖼️ Image question: Describe this image"


            # ------------------------------------------------
            # IF PDF EXISTS
            # ------------------------------------------------

            elif get_active_pdf():

                pdf = get_active_pdf()

                text_file = pdf["text_file"]

                text_path = os.path.join(
                    app.config["UPLOAD_FOLDER"],
                    text_file
                )


                if os.path.exists(
                    text_path
                ):

                    with open(
                        text_path,
                        "r",
                        encoding="utf-8"
                    ) as file:

                        full_text = file.read()


                    relevant_text = find_relevant_text(
                        full_text,
                        question
                    )


                    answer = ask_pdf_ai(
                        question,
                        relevant_text
                    )


                else:

                    answer = ask_ai(
                        question
                    )


            # ------------------------------------------------
            # NORMAL AI
            # ------------------------------------------------

            else:

                answer = ask_ai(
                    question
                )


            answer_html = markdown.markdown(
                answer
            )


            # Follow-up suggestions disabled (were slowing down
            # every answer by making a second AI call)

            follow_up_questions = []


            # SAVE HISTORY

            conn = get_db()

            cursor = conn.cursor()


            cursor.execute(
                """
                INSERT INTO chat_history(
                    username,
                    question,
                    answer,
                    image_file
                )
                VALUES(?,?,?,?)
                """,
                (
                    session["user"],
                    question,
                    answer_html,
                    image_saved_name
                )
            )


            conn.commit()

            last_chat_id = cursor.lastrowid

            conn.close()

            session["last_chat_id"] = last_chat_id

            session["last_image_file"] = image_saved_name

            session["last_follow_ups"] = follow_up_questions


        except Exception as e:

            print(
                "CHAT ERROR:",
                e
            )

            answer = (
                "Something went wrong: "
                + str(e)
            )

            session["last_follow_ups"] = []


    last_chat_id = session.get("last_chat_id")

    last_chat_rating = None

    last_chat_share_id = None


    if last_chat_id:

        conn = get_db()

        cursor = conn.cursor()

        cursor.execute(
            """
            SELECT rating, share_id
            FROM chat_history
            WHERE id=? AND username=?
            """,
            (last_chat_id, session["user"])
        )

        rating_row = cursor.fetchone()

        conn.close()


        if rating_row:

            last_chat_rating = rating_row["rating"]

            last_chat_share_id = rating_row["share_id"]


    return render_template(
        "index.html",
        username=session["user"],
        answer=(
            markdown.markdown(answer)
            if answer
            else ""
        ),
        question=question,
        pdfs=get_user_pdfs(),
        active_pdf_id=session.get("active_pdf_id"),
        last_chat_id=last_chat_id,
        last_image_file=session.get("last_image_file"),
        last_chat_rating=last_chat_rating,
        last_chat_share_id=last_chat_share_id,
        follow_up_questions=session.get("last_follow_ups", [])
    )


# =========================================================
# PDF UPLOAD
# =========================================================

@app.route(
    "/upload_pdf",
    methods=["POST"]
)
def upload_pdf():

    if "user" not in session:

        return redirect("/login")


    if "pdf" not in request.files:

        flash(
            "No PDF selected.",
            "error"
        )

        return redirect("/chat")


    file = request.files["pdf"]


    if file.filename == "":

        flash(
            "Please select a PDF.",
            "error"
        )

        return redirect("/chat")


    if not allowed_file(
        file.filename
    ):

        flash(
            "Only PDF files are allowed.",
            "error"
        )

        return redirect("/chat")


    username = secure_filename(
        session["user"]
    )


    original_name = secure_filename(
        file.filename
    )


    # Insert a placeholder DB row first so we get a unique id
    # (this id is used to build unique filenames, so multiple
    # PDFs with the same original name never collide)

    conn = get_db()

    cursor = conn.cursor()


    cursor.execute(
        """
        INSERT INTO pdf_documents(
            username,
            pdf_name,
            text_file,
            pdf_file
        )
        VALUES(?,?,?,?)
        """,
        (
            session["user"],
            original_name,
            "",
            ""
        )
    )


    conn.commit()

    new_pdf_id = cursor.lastrowid


    pdf_filename = (
        username
        + "_"
        + str(new_pdf_id)
        + "_"
        + original_name
    )


    pdf_path = os.path.join(
        app.config["UPLOAD_FOLDER"],
        pdf_filename
    )


    file.save(
        pdf_path
    )


    # Extract text / OCR

    extracted_text = extract_pdf_text(
        pdf_path
    )


    if not extracted_text:

        try:

            if os.path.exists(
                pdf_path
            ):

                os.remove(
                    pdf_path
                )

        except Exception:
            pass


        cursor.execute(
            """
            DELETE FROM pdf_documents
            WHERE id=?
            """,
            (new_pdf_id,)
        )

        conn.commit()

        conn.close()


        flash(
            "PDF uploaded, but no readable text was found. "
            "Please check your OCR installation.",
            "error"
        )

        return redirect("/chat")


    # Save extracted text

    text_filename = (
        username
        + "_"
        + str(new_pdf_id)
        + "_"
        + os.path.splitext(
            original_name
        )[0]
        + ".txt"
    )


    text_filename = secure_filename(
        text_filename
    )


    text_path = os.path.join(
        app.config["UPLOAD_FOLDER"],
        text_filename
    )


    with open(
        text_path,
        "w",
        encoding="utf-8"
    ) as text_file:

        text_file.write(
            extracted_text
        )


    # Update the DB row with the actual saved filenames

    cursor.execute(
        """
        UPDATE pdf_documents
        SET text_file=?, pdf_file=?
        WHERE id=?
        """,
        (
            text_filename,
            pdf_filename,
            new_pdf_id
        )
    )


    conn.commit()

    conn.close()


    # Make the newly uploaded PDF the active one

    session["active_pdf_id"] = new_pdf_id

    session["pdf_name"] = original_name


    flash(
        "PDF uploaded successfully! "
        "You can now ask questions about it.",
        "success"
    )


    return redirect("/chat")


# =========================================================
# REMOVE PDF
# =========================================================

@app.route(
    "/remove_pdf/<int:pdf_id>",
    methods=["POST"]
)
def remove_pdf(pdf_id):

    if "user" not in session:

        return redirect("/login")


    conn = get_db()

    cursor = conn.cursor()


    cursor.execute(
        """
        SELECT *
        FROM pdf_documents
        WHERE id=? AND username=?
        """,
        (pdf_id, session["user"])
    )

    pdf = cursor.fetchone()


    if not pdf:

        conn.close()

        flash(
            "PDF not found.",
            "error"
        )

        return redirect("/chat")


    if pdf["pdf_file"]:

        pdf_path = os.path.join(
            app.config["UPLOAD_FOLDER"],
            pdf["pdf_file"]
        )

        try:

            if os.path.exists(pdf_path):

                os.remove(pdf_path)

        except Exception as e:

            print(
                "PDF DELETE ERROR:",
                e
            )


    if pdf["text_file"]:

        text_path = os.path.join(
            app.config["UPLOAD_FOLDER"],
            pdf["text_file"]
        )

        try:

            if os.path.exists(text_path):

                os.remove(text_path)

        except Exception as e:

            print(
                "TEXT DELETE ERROR:",
                e
            )


    cursor.execute(
        """
        DELETE FROM pdf_documents
        WHERE id=?
        """,
        (pdf_id,)
    )


    conn.commit()

    conn.close()


    # If the removed PDF was the active one, clear the session

    if session.get("active_pdf_id") == pdf_id:

        session.pop("active_pdf_id", None)

        session.pop("pdf_name", None)


    flash(
        "PDF removed successfully.",
        "success"
    )


    return redirect("/chat")


# =========================================================
# SELECT ACTIVE PDF
# =========================================================

@app.route(
    "/select_pdf/<int:pdf_id>",
    methods=["POST"]
)
def select_pdf(pdf_id):

    if "user" not in session:

        return redirect("/login")


    conn = get_db()

    cursor = conn.cursor()


    cursor.execute(
        """
        SELECT *
        FROM pdf_documents
        WHERE id=? AND username=?
        """,
        (pdf_id, session["user"])
    )

    pdf = cursor.fetchone()

    conn.close()


    if not pdf:

        flash(
            "PDF not found.",
            "error"
        )

        return redirect("/chat")


    session["active_pdf_id"] = pdf_id

    session["pdf_name"] = pdf["pdf_name"]


    flash(
        f"Now chatting with: {pdf['pdf_name']}",
        "success"
    )


    return redirect("/chat")


# =========================================================
# HISTORY
# =========================================================

@app.route("/history")
def history():

    if "user" not in session:

        return redirect("/login")


    search_query = request.args.get(
        "q",
        ""
    ).strip()


    conn = get_db()

    cursor = conn.cursor()


    if search_query:

        like_pattern = "%" + search_query + "%"

        cursor.execute(
            """
            SELECT *
            FROM chat_history
            WHERE username=?
            AND (question LIKE ? OR answer LIKE ?)
            ORDER BY id DESC
            """,
            (
                session["user"],
                like_pattern,
                like_pattern
            )
        )

    else:

        cursor.execute(
            """
            SELECT *
            FROM chat_history
            WHERE username=?
            ORDER BY id DESC
            """,
            (session["user"],)
        )


    history_data = cursor.fetchall()

    conn.close()


    return render_template(
        "history.html",
        username=session["user"],
        history=history_data,
        search_query=search_query
    )


# =========================================================
# DELETE SINGLE CHAT
# =========================================================

@app.route(
    "/delete_chat/<int:chat_id>",
    methods=["POST"]
)
def delete_chat(chat_id):

    if "user" not in session:

        return redirect("/login")


    conn = get_db()

    cursor = conn.cursor()


    cursor.execute(
        """
        DELETE FROM chat_history
        WHERE id=? AND username=?
        """,
        (chat_id, session["user"])
    )


    conn.commit()

    conn.close()


    flash(
        "Chat deleted.",
        "success"
    )


    return redirect(request.referrer or "/history")


# =========================================================
# SHARE A CHAT (public read-only link)
# =========================================================

@app.route(
    "/share_chat/<int:chat_id>",
    methods=["POST"]
)
def share_chat(chat_id):

    if "user" not in session:

        return redirect("/login")


    conn = get_db()

    cursor = conn.cursor()


    cursor.execute(
        """
        SELECT *
        FROM chat_history
        WHERE id=? AND username=?
        """,
        (chat_id, session["user"])
    )

    row = cursor.fetchone()


    if not row:

        conn.close()

        flash(
            "Chat not found.",
            "error"
        )

        return redirect(request.referrer or "/chat")


    share_id = row["share_id"]


    if not share_id:

        share_id = secrets.token_urlsafe(8)

        cursor.execute(
            """
            UPDATE chat_history
            SET share_id=?
            WHERE id=?
            """,
            (share_id, chat_id)
        )

        conn.commit()


    conn.close()


    share_url = (
        request.host_url.rstrip("/")
        + "/share/"
        + share_id
    )


    flash(
        "🔗 Share link ready: " + share_url,
        "success"
    )


    return redirect(request.referrer or "/chat")


@app.route("/share/<token>")
def view_shared_chat(token):

    conn = get_db()

    cursor = conn.cursor()


    cursor.execute(
        """
        SELECT *
        FROM chat_history
        WHERE share_id=?
        """,
        (token,)
    )

    item = cursor.fetchone()

    conn.close()


    return render_template(
        "share.html",
        item=item
    )


# =========================================================
# RATE ANSWER (👍 / 👎)
# =========================================================

@app.route(
    "/rate_chat/<int:chat_id>/<rating_value>",
    methods=["POST"]
)
def rate_chat(chat_id, rating_value):

    if "user" not in session:

        return redirect("/login")


    try:

        rating_value = int(rating_value)

    except ValueError:

        flash(
            "Invalid rating.",
            "error"
        )

        return redirect(request.referrer or "/chat")


    # rating_value: 1 = like, -1 = dislike, 0 = clear

    if rating_value not in (1, -1, 0):

        flash(
            "Invalid rating.",
            "error"
        )

        return redirect(request.referrer or "/chat")


    conn = get_db()

    cursor = conn.cursor()


    cursor.execute(
        """
        SELECT rating
        FROM chat_history
        WHERE id=? AND username=?
        """,
        (chat_id, session["user"])
    )

    row = cursor.fetchone()


    if not row:

        conn.close()

        flash(
            "Chat not found.",
            "error"
        )

        return redirect(request.referrer or "/chat")


    # Clicking the same rating again clears it (toggle off)

    new_rating = (
        None
        if row["rating"] == rating_value
        else (
            rating_value
            if rating_value != 0
            else None
        )
    )


    cursor.execute(
        """
        UPDATE chat_history
        SET rating=?
        WHERE id=? AND username=?
        """,
        (new_rating, chat_id, session["user"])
    )


    conn.commit()

    conn.close()


    return redirect(request.referrer or "/chat")


# =========================================================
# EXPORT CHAT ANSWER (PDF / Word)
# =========================================================

@app.route("/export_chat/<int:chat_id>/<file_format>")
def export_chat(chat_id, file_format):

    if "user" not in session:

        return redirect("/login")


    conn = get_db()

    cursor = conn.cursor()


    cursor.execute(
        """
        SELECT *
        FROM chat_history
        WHERE id=? AND username=?
        """,
        (chat_id, session["user"])
    )

    chat_row = cursor.fetchone()

    conn.close()


    if not chat_row:

        flash(
            "Chat not found.",
            "error"
        )

        return redirect("/history")


    question = chat_row["question"]

    answer_html = chat_row["answer"]


    safe_name = re.sub(
        r"[^a-zA-Z0-9]+",
        "_",
        question[:40]
    ).strip("_") or "askai_answer"


    if file_format == "pdf":

        file_buffer = generate_pdf_file(
            question,
            answer_html
        )

        return send_file(
            file_buffer,
            as_attachment=True,
            download_name=safe_name + ".pdf",
            mimetype="application/pdf"
        )


    elif file_format == "docx":

        file_buffer = generate_docx_file(
            question,
            answer_html
        )

        return send_file(
            file_buffer,
            as_attachment=True,
            download_name=safe_name + ".docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


    flash(
        "Invalid export format.",
        "error"
    )

    return redirect("/history")


# =========================================================
# CLEAR HISTORY
# =========================================================

@app.route("/clear_history")
def clear_history():

    if "user" not in session:

        return redirect("/login")


    conn = get_db()

    cursor = conn.cursor()


    cursor.execute(
        """
        DELETE FROM chat_history
        WHERE username=?
        """,
        (session["user"],)
    )


    conn.commit()

    conn.close()


    flash(
        "Chat history cleared successfully.",
        "success"
    )


    return redirect("/history")


# =========================================================
# RUN
# =========================================================

if __name__ == "__main__":

    app.run(
        host="0.0.0.0",
        port=5000,
        debug=True
    )